from flask import Blueprint, request, jsonify, send_file
import os, re, json, pdfplumber
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO
from docx.shared import Pt
from Routing import create_anthropic_completion
from auth_utils import require_auth
cv_bp = Blueprint('cv', __name__)

# Anchor to this file's directory so paths match Assist.py regardless of cwd
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CV_DIR = os.path.join(BASE_DIR, "Cv_docs")
REFINED_CV_DIR = os.path.join(BASE_DIR, "refined")
CHAT = os.path.join(BASE_DIR, "chat")
os.makedirs(CV_DIR, exist_ok=True)
os.makedirs(REFINED_CV_DIR, exist_ok=True)
os.makedirs(CHAT, exist_ok=True)

def clean_text(text):
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()

def extract_pdf(path):
    with pdfplumber.open(path) as pdf:
        return clean_text("\n".join(
            p.extract_text() for p in pdf.pages if p.extract_text()
        ))

def extract_docx(path):
    doc = Document(path)
    return clean_text("\n".join(
        p.text for p in doc.paragraphs if p.text.strip()
    ))

def get_cv():
    # sort by modified time so the latest upload wins
    files = sorted(
        (f for f in os.listdir(CV_DIR) if f.endswith(".txt")),
        key=lambda f: os.path.getmtime(os.path.join(CV_DIR, f))
    )
    if not files:
        return ""
    with open(os.path.join(CV_DIR, files[-1]), "r") as f:
        return f.read().strip()

def get_relevant_history(keywords=("cv", "job", "skill", "project", "experience", "role", "apply")):
    chat_path = os.path.join(CHAT, "chat_history.txt")
    if not os.path.exists(chat_path):
        return []
    with open(chat_path, "r") as f:
        history = f.readlines()
    relevant = [
        line for line in history
        if any(keyword in line.lower() for keyword in keywords)
    ]
    return relevant[-10:]  # last 10 relevant messages max
def extract_anthropic_text(msg):
    if msg is None:
        return ""
    content = getattr(msg, "content", None)
    if isinstance(content, list):
        return "".join(
            getattr(block, "text", "") for block in content
            if getattr(block, "type", None) == "text"
        ).strip()
    return str(content or msg).strip()

@cv_bp.route('/cv/upload', methods=['POST'])
@require_auth
def upload():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file"}), 400

    filename = secure_filename(file.filename)
    temp = os.path.join(CV_DIR, filename)
    file.save(temp)

    if filename.endswith(".pdf"):
        text = extract_pdf(temp)
    elif filename.endswith((".docx", ".doc")):
        text = extract_docx(temp)
    else:
        os.remove(temp)
        return jsonify({"error": "PDF or DOCX only"}), 400

    os.remove(temp)
    saved = os.path.join(CV_DIR, filename.rsplit('.', 1)[0] + ".txt")
    with open(saved, "w") as f:
        f.write(text)

    return jsonify({"text": "CV uploaded", "cv":f"{CV_DIR}/{filename.rsplit('.', 1)[0]}.txt","preview": text[:300]})


@cv_bp.route('/cv/review', methods=['POST'])
@require_auth
def review():
    cv_text = get_cv()

    if not cv_text:
        return jsonify({"error": "No CV found"}), 400

    data = request.get_json(silent=True) or {}
    target = data.get("target_role", "tech/developer roles")

    try:
        prompt = f"""
        Review this CV for {target}. Return plain text only:

        {{
            \"score\": 1-10,
            \"strengths\": [\"max 3\"],
            \"weaknesses\": [\"max 3\"],
            \"missing\": [\"missing sections\"],
            \"summary\": \"2 sentence verdict\"
        }}

        CV:
        {cv_text[:3000]}
        """
        response = create_anthropic_completion(
            model="claude-sonnet-4-6",
            system="You are a helpful assistant, help the user review their CV.",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000
        )
        response_text = extract_anthropic_text(response)

        match = re.search(r"\{.*\}", response_text, re.DOTALL)
        if not match:
            raise ValueError("No JSON found")
        review_data = json.loads(match.group())
        cv_filename = os.path.join(REFINED_CV_DIR, f"review_{target.replace(' ', '_')}.json")

        with open(cv_filename, "w") as f:
            json.dump(review_data, f, indent=2)
        n_review = {
            "score": review_data.get("score", ""),
            "strengths": review_data.get("strengths", []),
            "weaknesses": review_data.get("weaknesses", []),
            "missing": review_data.get("missing", []),
            "summary": review_data.get("summary", "")
        }
        return jsonify({"review": n_review})

    except json.JSONDecodeError:
        return jsonify({"error": "Model returned invalid JSON"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@cv_bp.route('/cv/rewrite', methods=['POST'])
@require_auth
def rewrite():
    cv_text = get_cv()
    if not cv_text:
        return jsonify({"error": "No CV found"}), 400
    data = request.get_json(silent=True) or {}
    target = data.get("target_role", "tech/developer roles")
    feedback = os.path.join(REFINED_CV_DIR, f"review_{target.replace(' ', '_')}.json")
    review_data = {}
    if os.path.exists(feedback):
        with open(feedback, "r") as f:
            review_data = json.load(f)
    chat_history = get_relevant_history()
    try:
        user_context = """
                        About the user:
                        - Name: Godwin
                        - Age: 21, based in Horley, Surrey
                        - Currently studying HNC/HND Computing
                        - Progressing to a one year Computer Science top up degree (University of Greater Manchester, 2027)
                        - Works in fast food currently
                        - Building Twinssistant and Alamu — a multi provider AI assistant with voice, memory, and agentic routing
                        - Targeting junior backend developer and AI integration roles
                        - Skills: Python, React, Flask, Firebase, Firestore, REST APIs, WebSockets, Anthropic/OpenAI/Gemini APIs
                        """
        prompt = f"""
                    Rewrite this CV for {target}.
                    Use the userdetails{user_context},and {review_data} and {chat_history} as feedback to improve the CV. Focus on addressing weaknesses and missing sections, while maintaining strengths.
                        - Keep all real experience
                        - Achievement-focused bullet points
                        - Strong action verbs
                        - Remove filler
                        - Do not add any em-dashes or emojis
                        - Tailor skills and keywords to the target role
                        - Take what the user has and expand on it, do not remove content unless it's clearly weak or irrelevant
                        - If experience is very thin, creatively expand on it to make it more substantial, but do not fabricate any new experience. For example, if they have a one-line experience about "built a website", you could expand that into multiple bullet points about the technologies used, the features of the website, the impact it had, etc. Just be creative with how you describe and expand on the existing experience.
                        - If a section is missing, add it with placeholder content based on what you know about the user from the CV and the target role. For example, if they have no "Skills" section but mention various technologies in their experience, you could create a "Skills" section that lists those technologies.
                        - Use concise language, avoid fluff
                        - Use the best cv examples online as reference, but do not copy any phrasing or formatting, make it original
                        Return plain text only, ready to copy.

                        CV:
                        {cv_text[:3000]}"""
        response = create_anthropic_completion(
                    model="claude-sonnet-4-6",
                    messages=[
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=2048,
                    temperature=0.7
                )
        response_text = extract_anthropic_text(response)
        doc = Document()

        # set font defaults
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        for line in response_text.splitlines():
            if not line.strip():
                doc.add_paragraph()  # blank spacer
                continue

            # strip markdown bold/italic
            clean = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', line).strip()
            # strip markdown headers
            is_heading = clean.startswith('#')
            clean = clean.lstrip('#').strip()

            # detect bullet points
            is_bullet = line.strip().startswith(('-', '*', '•'))
            if is_bullet:
                clean = re.sub(r'^[-*•]\s*', '', clean)

            if is_heading or (clean.isupper() and len(clean) < 50):
                p = doc.add_paragraph()
                run = p.add_run(clean)
                run.bold = True
                run.font.size = Pt(13)
                p.paragraph_format.space_before = Pt(10)
            elif is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(clean)
            else:
                p = doc.add_paragraph()
                p.add_run(clean)
        buffer = BytesIO()
        doc.save(buffer)
        with open(os.path.join(REFINED_CV_DIR, f"rewritten_{target.replace(' ', '_')}.docx"), "wb") as f:
            f.write(buffer.getvalue())
        buffer.seek(0)
        return send_file(buffer,as_attachment=True,
                         download_name=f"rewritten_{target.replace(' ', '_')}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return jsonify({"error": str(e)}), 500